from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "seed-cup-embedded-dual-adc-task.md"
OUTPUT = ROOT / "docs" / "seed-cup-embedded-dual-adc-task.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_left_border(paragraph, color=BLUE, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def inline_runs(paragraph, text, size=11, color="000000"):
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=size - 0.3, color=INK)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_body_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    inline_runs(paragraph, text)
    return paragraph


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(paragraph, LIGHT_GRAY)
    set_left_border(paragraph, color="9AA8B8", size="10", space="6")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.2, color=INK)
        if index < len(lines) - 1:
            run.add_break()
    return paragraph


def add_table(doc, rows):
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    if column_count == 4:
        widths = [1500, 2300, 2100, 3460]
    elif column_count == 3:
        widths = [1600, 1100, 6660]
    else:
        widths = [4680, 4680]
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            inline_runs(paragraph, value, size=9.4 if row_index else 9.6, color=INK if row_index else DARK_BLUE)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.7)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    inline_runs(header, "种子杯 · 嵌入式赛道", size=9, color=MUTED)
    right_run = header.add_run("    AliceSIM 题目发布稿")
    set_run_font(right_run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    inline_runs(footer, "AliceSIM · 2026    第 ", size=9, color=MUTED)
    add_page_field(footer)

    index = 0
    first_content = True
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, size=23, color=INK, bold=True)
            first_content = False
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
        elif stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            if rows:
                add_table(doc, rows)
            continue
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            inline_runs(p, stripped[2:].strip(), size=10.7)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped), size=10.7)
        else:
            p = add_body_paragraph(doc, stripped)
            if first_content:
                set_paragraph_shading(p, CALLOUT)
                set_left_border(p)
                p.paragraph_format.left_indent = Inches(0.12)
                p.paragraph_format.right_indent = Inches(0.12)
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(10)
                first_content = False
        index += 1

    doc.core_properties.title = "2026“种子杯”嵌入式赛道题目：双路模拟量采集与状态指示终端"
    doc.core_properties.subject = "AliceSIM STM32F103C8T6 embedded track problem"
    doc.core_properties.author = "AliceSIM"
    doc.core_properties.keywords = "种子杯, AliceSIM, STM32F103C8T6, ADC, OLED, UART"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
